# -*- coding: utf-8 -*-
"""Génère le dossier de projet QResto au format .docx"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
GREY = RGBColor(0x6B, 0x72, 0x80)
DARK = RGBColor(0x1F, 0x29, 0x37)

doc = Document()

# ------------------------------------------------------------------ styles
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = sec.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)

for name, size, color, bold in (('Heading 1', 16, ACCENT, True),
                                ('Heading 2', 12.5, DARK, True),
                                ('Title', 28, ACCENT, True)):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = bold


def para(text='', *, bold=False, italic=False, size=10.5, color=None,
         align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p


def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def puce(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def shade(cell, hexa):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hexa)
    cell._tc.get_or_add_tcPr().append(el)


def table(entetes, lignes, largeurs):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(e)
        r.bold = True
        r.font.size = Pt(9.5)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        shade(c, 'F3F4F6')

    for ligne in lignes:
        cells = t.add_row().cells
        for i, val in enumerate(ligne):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)

    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(largeurs[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


C = WD_ALIGN_PARAGRAPH.CENTER

# --------------------------------------------------------------- couverture
for _ in range(5):
    doc.add_paragraph()
para('QResto', bold=True, size=40, color=ACCENT, align=C, space_after=4)
para('Commande par QR code pour les restaurants algériens',
     size=14, color=GREY, align=C, space_after=30)
para('Dossier de projet', bold=True, size=12, align=C, space_after=2)
para('Analyse, conception, réalisation', size=10.5, color=GREY, align=C, space_after=70)
para('Abdelkarim Laabani', bold=True, size=11, align=C, space_after=2)
para('Août 2026', size=10, color=GREY, align=C, space_after=2)
para('github.com/Karimmm06-lb/QResto', size=9.5, color=GREY, align=C)
doc.add_page_break()

# ------------------------------------------------------------- 1. contexte
h1('1. Contexte et problème')
para("Dans de nombreux restaurants algériens, la prise de commande reste entièrement "
     "orale. Aux heures de pointe, un serveur enchaîne quatre à cinq tables de suite et "
     "doit mémoriser qui a commandé quoi.")
para('Trois conséquences sont systématiquement observées :')
puce("des erreurs d'attribution : un plat servi à la mauvaise table ;")
puce("une attente du client avant même de pouvoir commander ;")
puce("aucune trace écrite des ventes en fin de service.")
para("Ces établissements sont peu équipés en informatique : leur présence numérique se "
     "limite le plus souvent à une page Facebook. Les solutions existantes — caisses "
     "tactiles, bornes de commande — sont conçues pour des chaînes et supposent un budget "
     "et une compétence technique qu'ils n'ont pas.")

h2('Établissements ciblés — Aïn Benian, Alger')
table(['Établissement', 'Type', 'Présence en ligne'],
      [['Black & Silver', 'Pizza, burgers, poulet frit', 'Page Facebook'],
       ['Spicymax', 'Pizza au feu de bois, burgers', 'Page Facebook'],
       ['Pizza Home', 'Pizza, tacos, sandwichs', 'Page Facebook'],
       ["Pizzeria L'Abri-Côtier", 'Pizza, pâtes', 'Page Facebook'],
       ['Team Pizza', 'Pizzeria', 'Fiche Tripadvisor'],
       ['Khayma', 'Cuisine traditionnelle', 'Page Facebook']],
      [5.0, 6.0, 4.5])
para("Les fast-foods à forte rotation sont les meilleurs candidats : c'est là que le "
     "serveur sature et que le gain est le plus visible.", italic=True, color=GREY)

# ------------------------------------------------------------- 2. solution
h1('2. La solution')
table(['Étape', 'Acteur', 'Action'],
      [['1', 'Client', 'Scanne le QR code de sa table, le menu apparaît sur son téléphone'],
       ['2', 'Client', 'Compose sa commande, indique son prénom, valide'],
       ['3', 'Caissier', 'La commande apparaît en direct sur son écran, avec un bip'],
       ['4', 'Caissier', 'Imprime le ticket et le remet au cuisinier'],
       ['5', 'Cuisinier', "Prépare — il ne change rien à ses habitudes"],
       ['6', 'Serveur', 'Livre le plat, le prénom lui indique le destinataire'],
       ['7', 'Client', 'Paie la totalité de sa table en caisse avant de partir']],
      [1.6, 2.4, 11.5])

para('Trois exclusions volontaires', bold=True)
puce("Pas de paiement en ligne : le paiement électronique est peu répandu en Algérie et "
     "les restaurateurs tiennent à encaisser eux-mêmes.")
puce("Pas de compte client : aucune inscription, aucune application à installer. C'est un "
     "argument commercial majeur.")
puce("Pas de notification au serveur : l'objectif est de le décharger, pas de lui ajouter "
     "du bruit.")

# ---------------------------------------------------------------- 3. cycle
h1('3. Cycle de développement')
para("Le projet suit les six phases du cycle de développement logiciel. Une particularité "
     "méthodologique doit être signalée : un prototype fonctionnel a été réalisé avant la "
     "formalisation des phases 1 à 3.")
para("Ce choix visait à valider l'expérience utilisateur et la faisabilité du temps réel "
     "avant d'investir dans l'infrastructure. Il a eu un coût réel : les décisions de "
     "cadrage prises ensuite ont conduit à revoir le modèle de données, en particulier "
     "l'introduction de la session de table, absente du prototype.")
table(['Phase', 'État', 'Livrable'],
      [['1. Planification', 'Terminée', 'Périmètre, faisabilité, six risques, critères de succès'],
       ['2. Analyse', 'Terminée', '22 besoins fonctionnels, 10 non fonctionnels, 14 règles de gestion'],
       ['3. Conception', 'Terminée', 'Architecture, modèle de données, contrats, 22 décisions'],
       ['4. Implémentation', 'Terminée',
        'Quatre interfaces, base déployée en production, deux cartes réelles importées'],
       ['5. Tests et intégration', 'En cours',
        "Tests d'intégration et de bout en bout passés ; recette sur site à mener"],
       ['6. Maintenance', 'À venir', 'Exploitation, supervision, évolutions']],
      [3.8, 2.4, 9.3])

para("Vingt-trois décisions ont été prises à ce jour, dont une correction : D5-bis a "
     "rouvert D5 après que la carte réelle d'un restaurant cible eut invalidé l'hypothèse "
     "sur laquelle elle reposait. Sept décisions restent ouvertes ; trois d'entre elles "
     "attendent des observations de terrain.")

# ------------------------------------------------------------ 4. décisions
h1('4. Décisions structurantes')
para("Vingt-deux décisions ont été identifiées et tranchées avant l'implémentation. "
     "Chacune est consignée avec sa justification et les alternatives écartées. Les plus "
     "déterminantes sont reprises ci-dessous.")
table(['Réf.', 'Décision retenue', 'Justification'],
      [['D1', 'Session de table comme unité de facturation',
        "Un groupe commande en plusieurs vagues mais règle une seule addition. Sans session, "
        "le caissier doit deviner quelles commandes vont ensemble."],
       ['D2', 'Exécution directe, sans validation humaine',
        "Imposer une validation déplacerait le goulot d'étranglement du serveur vers le "
        "caissier : le problème d'origine ne serait pas résolu. L'impression vaut engagement."],
       ['D3a', 'Jeton statique par table, détection humaine',
        "Le caissier voit la salle. Une table affichée sans clients signale une commande "
        "frauduleuse. Les mécanismes techniques ajoutent une friction permanente pour un "
        "risque non avéré."],
       ['D3b', 'Prénom du convive sur la commande',
        "Permet au serveur de remettre la bonne assiette à la bonne personne. Donnée "
        "déclarative, jamais un moyen de contrôle."],
       ['D4', 'Multi-tenant dès la première migration',
        "Le surcoût immédiat est de quelques dizaines de lignes. Le rattrapage ultérieur "
        "imposerait de réécrire la sécurité sur une base en production."],
       ['D5', 'Déclinaisons de plats obligatoires',
        "Certains restaurants ont des tailles, d'autres non : le schéma couvre le "
        "sur-ensemble. Un plat simple reçoit une déclinaison masquée par l'interface."],
       ['D5-bis', 'Suppléments rattachés à la ligne de commande',
        "Corrige D5. Les cartes réelles ont invalidé l'hypothèse « aucun restaurant cible "
        "n'a exprimé ce besoin » : les suppléments y sont partout, et sur les pizzas leur "
        "prix dépend de la taille. Traités comme des articles séparés, ils rendaient le "
        "ticket cuisine ambigu."],
       ['D6', 'Disponibilité basculée manuellement',
        "Un décrément automatique suppose un inventaire tenu à jour au produit près, "
        "qu'aucun établissement cible ne tient. Un stock faux est pire qu'aucun stock."],
       ['D7', 'Annulation par la caisse, motif après impression',
        "L'annulation est le seul rattrapage depuis que D2 a supprimé la validation. Après "
        "impression, la perte est réelle et doit être tracée."],
       ['D8', "Temps d'attente indicatif, en fourchette",
        "Un chiffre précis et faux crée du conflit ; une fourchette honnête crée de la confiance."],
       ['D21', 'Expiration à 4 h et clôture de journée',
        "Sans mécanisme, les sessions orphelines faussent les statistiques. La journée "
        "d'exploitation se termine à 4 h, pas à minuit."],
       ['D22', 'Suivi client par interrogation périodique',
        "Le client anonyme n'a aucun droit de lecture sur les commandes. Lui ouvrir un accès "
        "permettrait à un concurrent de compter les commandes du restaurant."]],
      [1.5, 5.0, 9.0])

# --------------------------------------------------------- 5. architecture
h1('5. Architecture')
para('Architecture Jamstack multi-tenant, à trois niveaux, sans serveur applicatif, '
     'pilotée par les événements.', bold=True)
table(['Terme', 'Signification dans ce projet'],
      [['Jamstack', "Le front est statique et servi par un réseau de diffusion ; tout le "
                    "dynamique passe par des appels d'API depuis le navigateur."],
       ['Trois niveaux', "Présentation (les pages) · Logique métier (PostgreSQL) · "
                         "Données (PostgreSQL)."],
       ['Sans serveur applicatif', "Aucun serveur à écrire ni à administrer : la base "
                                   "managée expose directement l'API."],
       ['Piloté par les événements', "Le poste caisse s'abonne et reçoit une notification "
                                     "poussée ; il n'interroge pas en boucle."],
       ['Multi-tenant', "Une instance unique sert tous les restaurants, cloisonnés par "
                        "leur identifiant."]],
      [4.5, 11.0])
para("Ce choix est imposé par une contrainte : un restaurant ne peut ni payer d'hébergement "
     "ni administrer un serveur. Cette architecture a un coût d'exploitation nul et aucune "
     "administration système.")

h2('Asymétrie de diffusion')
para("Le poste caisse et le téléphone du client n'utilisent pas le même mécanisme, et c'est "
     "délibéré.")
table(['Destinataire', 'Mécanisme', 'Exigence de latence'],
      [['Poste caisse', 'Abonnement, notification poussée', 'Moins de 3 secondes'],
       ['Téléphone du client', 'Interrogation toutes les 10 secondes', 'Aucune exigence formelle']],
      [4.5, 6.5, 4.5])

h2('Technologies')
table(['Couche', 'Choix', 'Motif'],
      [['Interfaces', 'HTML, CSS et JavaScript natif',
        'Aucune dépendance, aucune compilation, hébergement statique gratuit'],
       ['Base de données', 'PostgreSQL (Supabase)',
        'Logique métier et sécurité au plus près des données'],
       ['Temps réel', 'Diffusion par WebSocket',
        'Notification poussée respectant les politiques de sécurité'],
       ['Authentification', 'Comptes gérés par le service',
        'Le restaurant est porté par le jeton, non falsifiable'],
       ['Impression', "Moteur d'impression du navigateur",
        'Une imprimante thermique est vue comme une imprimante ordinaire']],
      [3.2, 4.8, 7.5])

# ------------------------------------------------------------- 6. données
h1('6. Modèle de données')
para("Onze tables, trente et un index, onze politiques de sécurité, quinze procédures et "
     "un déclencheur. Les éléments qui suivent sont extraits du schéma réellement déployé.")

h2('Modèle conceptuel')
para('La chaîne principale :')
para('RESTAURANT  →  TABLE  →  SESSION  →  COMMANDE  →  LIGNE DE COMMANDE',
     bold=True, align=C, space_after=10)
table(['Entité', 'Définition métier'],
      [['Restaurant', 'Établissement abonné au service, avec ses paramètres de fonctionnement'],
       ['Table', 'Emplacement physique dans la salle, porteur du QR code'],
       ['Catégorie', 'Regroupement de plats dans la carte'],
       ['Plat', 'Article de la carte. Un supplément est un plat marqué comme tel'],
       ['Variante', 'Déclinaison d\'un plat portant le prix'],
       ['Session', 'Le couvert : un groupe installé à une table, unité de facturation'],
       ['Commande', 'Un envoi effectué par un convive pendant une session'],
       ['Ligne de commande', 'Une variante et sa quantité dans une commande']],
      [3.5, 12.0])

para('Trois associations méritent attention', bold=True)
puce("SESSION s'intercale entre TABLE et COMMANDE. C'est le cœur du modèle : sans elle, "
     "trois commandes d'une même table seraient trois objets sans lien, et le caissier "
     "devrait deviner quoi encaisser.")
puce("LIGNE DE COMMANDE est en relation réflexive avec elle-même : une ligne « camembert » "
     "pointe vers la ligne « burger » qu'elle complète. Le ticket cuisine peut ainsi "
     "indiquer sur quel plat poser le supplément.")
puce("PLAT et CATÉGORIE sont liés en plusieurs-à-plusieurs pour porter la portée d'un "
     "supplément : les fromages s'appliquent aux burgers et aux sandwichs, les garnitures "
     "sucrées aux crêpes.")

h2('Modèle logique')
table(['Table', 'Colonnes principales'],
      [['restaurants', 'id, nom, ville, fuseau, fin_journee, validation_requise, code_table_requis, eta_active'],
       ['tables_resto', 'id, restaurant_id, numero, qr_token, active'],
       ['categories', 'id, restaurant_id, nom_fr, nom_ar, nom_en, ordre'],
       ['plats', 'id, restaurant_id, categorie_id, noms et descriptions trilingues, image_url, disponible, archive, est_supplement'],
       ['variantes_plat', 'id, plat_id, libelles trilingues, prix, ordre'],
       ['supplements_categories', 'supplement_id, categorie_id'],
       ['sessions', 'id, restaurant_id, table_id, statut, journee, total, ouverte_le, activite_le, fermee_le'],
       ['commandes', 'id, restaurant_id, session_id, numero, journee, statut, nom_convive, note, total, eta_min, eta_max, secret, imprimee_le, motif_annulation'],
       ['lignes_commande', 'id, commande_id, variante_id, plat_id, quantite, prix_unitaire, libelle, parent_ligne_id'],
       ['compteurs_journee', 'restaurant_id, journee, dernier_numero'],
       ['journal_audit', 'id, restaurant_id, commande_id, session_id, action, detail, acteur, cree_le']],
      [3.5, 12.0])

h2('Colonnes dont le choix mérite justification')
table(['Colonne', 'Pourquoi elle existe sous cette forme'],
      [['restaurants.fin_journee',
        "La journée d'exploitation se termine à 4 h, pas à minuit : un restaurant ouvert "
        "jusqu'à 1 h verrait sinon son service coupé en deux."],
       ['tables_resto.qr_token',
        "Le QR encode un identifiant aléatoire, jamais le numéro de table : un numéro "
        "serait devinable."],
       ['plats.archive',
        "Un plat retiré de la carte n'est jamais supprimé : les commandes passées le "
        "référencent encore."],
       ['plats.est_supplement',
        "Un supplément est un plat marqué comme tel : il hérite des déclinaisons, de la "
        "disponibilité et du prix figé sans table dédiée."],
       ['variantes_plat.prix',
        "Le prix vit sur la déclinaison, jamais sur le plat. Garder les deux créerait deux "
        "chemins de lecture et des incohérences."],
       ['commandes.secret',
        "Seul moyen pour le client de suivre sa commande sans exposer celles des autres "
        "tables."],
       ['commandes.imprimee_le',
        "L'impression du ticket est l'acte d'engagement de production."],
       ['lignes_commande.prix_unitaire',
        "Prix figé à la vente. Ce n'est pas une copie du prix courant : sans lui, une "
        "hausse de tarif fausserait rétroactivement tout l'historique."],
       ['lignes_commande.libelle',
        "Instantané du libellé. Un ticket réimprimé doit rester identique à l'original, "
        "même si le plat a été renommé depuis."],
       ['compteurs_journee',
        "Sans compteur dédié, deux commandes simultanées obtiendraient le même numéro."]],
      [4.5, 11.0])

h2('Contraintes portées par la base')
para("Ces règles ne sont pas écrites dans le navigateur. Le client n'étant pas authentifié, "
     "tout contrôle placé côté interface serait contournable.")
table(['Contrainte', 'Ce qu\'elle garantit'],
      [['Index unique partiel sur les sessions ouvertes',
        "Une table n'a qu'une seule session ouverte : deux convives tombent forcément sur "
        "la même addition."],
       ['Unicité du numéro par restaurant et par journée',
        'Pas de doublon de numéro de commande dans une journée.'],
       ['Contraintes de domaine sur les statuts',
        'Aucun état hors des machines à états définies.'],
       ['Déclencheur sur les suppléments',
        "Un supplément ne se commande pas seul, un plat ne se rattache pas à un plat, et "
        "un supplément n'en porte pas un autre."],
       ['Suppressions restreintes',
        "On ne supprime pas une table qui a servi, ni une déclinaison déjà commandée."]],
      [5.5, 10.0])

h2('Normalisation')
para("Le schéma est en troisième forme normale, avec trois dénormalisations assumées.")
table(['Champ', 'Justification'],
      [['sessions.total, commandes.total',
        "Le poste caisse affiche des montants en continu ; les recalculer à chaque "
        "rafraîchissement serait coûteux et inutile."],
       ['lignes_commande.plat_id',
        'Déductible via la déclinaison, mais les statistiques par plat exigeraient sinon '
        'une jointure de plus sur la table la plus volumineuse.'],
       ['lignes_commande.libelle',
        "Ce n'est pas une duplication mais un instantané : le ticket doit rester identique "
        "même si le plat est renommé."]],
      [4.5, 11.0])

h2('Volumétrie')
para("Moins de cinq cent mille lignes par an et par restaurant, dont l'essentiel en lignes "
     "de commande et en journal d'audit. Aucune contrainte de volumétrie : le plan gratuit "
     "couvre plusieurs années d'exploitation pour plusieurs établissements.")

# ------------------------------------------------------------ 7. sécurité
h1('7. Sécurité')
para("Principe directeur : le navigateur du client n'est jamais une source de confiance.",
     bold=True)
para('Deux conséquences structurent tout le système.')
puce("Aucun prix ne transite depuis le client. La procédure de création reçoit uniquement "
     "des identifiants de déclinaisons et des quantités ; les tarifs sont lus en base et le "
     "total y est recalculé. Modifier le JavaScript de la page n'a aucun effet.")
puce("Aucune écriture directe n'est autorisée. Les tables de commandes n'ont aucune "
     "politique d'insertion : le seul chemin d'écriture passe par des procédures stockées.")
table(['Ressource', 'Client anonyme', 'Caisse', 'Gérant'],
      [['Menu', 'Lecture', 'Lecture', 'Lecture et écriture'],
       ['Créer une commande', 'Procédure', '—', '—'],
       ['Suivre sa commande', 'Par secret', '—', '—'],
       ['Commandes du restaurant', 'Aucun accès', 'Lecture', 'Lecture'],
       ['Imprimer, annuler, encaisser', 'Aucun accès', 'Procédure', 'Procédure'],
       ["Journal d'audit", 'Aucun accès', 'Lecture', 'Lecture']],
      [5.0, 3.5, 3.0, 4.0])

h2("Failles détectées par l'audit et corrigées")
para("Un audit de sécurité a été passé sur la base après déploiement. Il a révélé deux "
     "défauts réels, tous deux introduits pendant la conception.")
para('1. Clôture de journée accessible à tous — grave', bold=True, space_after=3)
para("Le contrôle d'accès reposait sur la comparaison entre le restaurant demandé et celui "
     "du jeton. Pour un appelant anonyme, ce dernier vaut NULL : la comparaison vaut NULL, "
     "et une condition NULL est fausse. Le garde-fou ne se déclenchait donc jamais, et "
     "n'importe qui pouvait clôturer la journée de n'importe quel restaurant. C'est le piège "
     "classique du NULL en SQL : une comparaison avec NULL n'est ni vraie ni fausse. Corrigé "
     "par un test explicite, puis l'exploitation a été rejouée en rôle anonyme pour vérifier "
     "qu'elle échoue.")
para('2. Fonction de maintenance exposée — moyenne', bold=True, space_after=3)
para("La fonction d'expiration des sessions n'a aucun cloisonnement par restaurant. Appelée "
     "par un tiers, elle expirait les sessions de tous les restaurants. Les droits ont été "
     "retirés à tous les rôles.")
para("Cause commune : retirer les droits du rôle public ne suffit pas ; il faut les révoquer "
     "explicitement des rôles anonyme et authentifié.", italic=True, color=GREY)

# --------------------------------------------------------------- 8. tests
h1('8. Réalisation et tests')
h2("Tests d'intégration sur la base")
table(['Règle vérifiée', 'Scénario', 'Résultat'],
      [['Une seule addition par table', 'Deux convives commandent séparément',
        'Une session, addition de 4 500 DA'],
       ['Total calculé en base', 'Panier de trois lignes',
        'Montant exact, aucun prix reçu du client'],
       ['Jeton de table inconnu', 'Identifiant inventé', 'Rejeté'],
       ['Panier vide', 'Envoi sans article', 'Rejeté'],
       ['Quantité aberrante', 'Valeurs négative et à 999', 'Rejeté'],
       ['Plat épuisé dans un panier', 'Panier mixte, un plat indisponible',
        'Commande entièrement refusée'],
       ["Contrôle d'accès à la clôture", 'Appel en rôle anonyme', 'Rejeté après correctif'],
       ['Supplément commandé seul', 'Camembert sans plat', 'Rejeté'],
       ['Plat servant de supplément', 'Burger rattaché à un burger', 'Rejeté']],
      [4.5, 5.5, 5.5])

h2('Test de bout en bout dans le navigateur')
para("Le parcours complet a été déroulé : scan du QR, composition du panier avec "
     "déclinaisons, saisie du prénom et d'une remarque, envoi. La commande est apparue sur "
     "le poste caisse sans rafraîchissement, groupée par table. L'impression a fait passer "
     "la commande en cuisine, et le téléphone du client a reflété le changement. Le cycle "
     "s'est achevé par la mise en statut prête, puis servie, puis l'encaissement de la table "
     "pour 4 050 DA.")

h2('Mise en situation réelle')
para("Deux cartes de restaurants d'Aïn Benian ont été relevées et importées : Black & "
     "Silver, soixante-douze plats sur treize catégories, et Spicy Max, trente-huit plats "
     "sur sept catégories. Cette confrontation au réel a produit trois enseignements.")
puce("Elle a invalidé une décision : les suppléments, écartés au cadrage faute de besoin "
     "exprimé, sont présents sur les deux cartes. D5-bis a corrigé le modèle.")
puce("Elle a validé les déclinaisons : pizzas en deux tailles, tacos en deux formats, "
     "burgers à prix unique — les trois cas coexistent sur une même carte.")
puce("Elle a révélé le seul écart avec le cadrage encore ouvert : les photos des plats, "
     "annoncées comme l'une des quatre fonctionnalités retenues, ne sont pas livrées. Les "
     "sources publiques ne fournissent que quelques visuels basse définition, insuffisants "
     "et dont la republication demanderait l'accord des établissements.")

h2("Enseignement sur l'architecture")
para("La couche d'accès aux données avait été isolée dès le prototype, dans l'idée que "
     "passer de la mémoire du navigateur à une vraie base ne toucherait qu'un seul fichier. "
     "Le résultat est nuancé : la structure a tenu, mais le passage d'un accès synchrone à "
     "un accès asynchrone a imposé d'adapter les trois interfaces. L'abstraction protégeait "
     "de l'origine des données, pas de leur temporalité.")

# ---------------------------------------------------------------- 9. suite
h1('9. Ce qui reste à faire')
table(['Réf.', 'Sujet', 'Enjeu'],
      [['D10', 'Idempotence des envois',
        "Sur réseau instable, un double envoi crée aujourd'hui deux commandes"],
       ['D12', 'Coupure internet en service', 'Aucun mode dégradé prévu'],
       ['D13', "Panne d'imprimante", 'Aucun repli prévu'],
       ['D14', 'Saisie initiale du menu',
        'Risque le plus probable du projet, et il est organisationnel'],
       ['D16', 'Rétention et purge du prénom',
        'Le prénom doit être purgé à la clôture, les montants conservés'],
       ['D18', 'Sur place ou à emporter', 'Non traité'],
       ['D19', 'Modèle économique', 'Non traité'],
       ['D20', 'Limitation de débit', 'Protection contre le spam de commandes'],
       ['—', 'Photos des plats',
        "Fonctionnalité annoncée au cadrage, non livrée : les sources publiques sont "
        "insuffisantes et leur republication demanderait l'accord des établissements"]],
      [1.5, 5.0, 9.0])

para("Trois de ces décisions dépendent d'observations de terrain", bold=True)
para("D12, D13 et D14 portent sur la coupure d'internet en service, la panne d'imprimante "
     "et la maintenance du menu. Les trancher depuis un bureau produirait des hypothèses, "
     "pas des décisions. Une fiche d'observation a été préparée pour la première visite, "
     "où chaque ligne relevée alimente une décision précise.")

para("Le risque principal n'est pas technique", bold=True)
para("Le risque le plus probable du projet est la saisie initiale du menu : plusieurs heures "
     "de travail pour un restaurateur peu à l'aise avec l'informatique. Aucune ligne de code "
     "ne le résout. La réponse est organisationnelle : la saisie doit être assurée à "
     "l'installation, pas laissée au client.")

h2('Critères de succès')
para('Le projet sera considéré comme réussi si, sur un restaurant pilote :')
puce('une commande apparaît en caisse en moins de trois secondes ;')
puce("le personnel n'a reçu aucune formation au-delà d'une démonstration de dix minutes ;")
puce("le restaurant réalise une journée complète de service sans intervention extérieure ;")
puce('le gérant consulte spontanément ses statistiques après le service.')
para("Le dernier critère est le plus révélateur : il mesure l'usage réel, et non la "
     "conformité technique.", italic=True, color=GREY)

doc.save(sys.argv[1])
print('OK', sys.argv[1])
